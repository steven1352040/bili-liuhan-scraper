"""将 liuhan_analysis.md 转换为 Word 文档"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 处理中文字体
from docx.oxml.ns import qn
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

with open('output/liuhan_analysis.md', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if line.startswith('# ') and not line.startswith('## '):
        # H1 title
        p = doc.add_heading(line[2:], level=0)
    elif line.startswith('#### '):
        # H4 - episode heading
        title = line[5:]
        p = doc.add_heading(title, level=4)
    elif line.startswith('### '):
        # H3
        title = line[4:]
        p = doc.add_heading(title, level=3)
    elif line.startswith('## '):
        # H2
        title = line[3:]
        p = doc.add_heading(title, level=2)
    elif line.startswith('> '):
        # Blockquote
        p = doc.add_paragraph(line[2:])
        p.style = doc.styles['Quote'] if 'Quote' in [s.name for s in doc.styles] else 'Normal'
    elif line.startswith('| ') and '|' in line:
        # Table handling
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            rows.append(row)
            i += 1
        i -= 1  # will be incremented

        if len(rows) >= 2:
            # Skip separator row (---|---)
            header = rows[0]
            data_rows = [r for r in rows[1:] if not all(c.replace('-','').strip() == '' for c in r)]

            table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
            table.style = 'Light Grid Accent 1'

            for j, h in enumerate(header):
                table.rows[0].cells[j].text = h
                for p in table.rows[0].cells[j].paragraphs:
                    for run in p.runs:
                        run.bold = True

            for ri, row in enumerate(data_rows):
                for ci, cell in enumerate(row):
                    if ci < len(header):
                        table.rows[ri+1].cells[ci].text = cell

            doc.add_paragraph()  # spacing after table
    elif line.startswith('- ') or line.startswith('  - '):
        p = doc.add_paragraph(line.strip('- '), style='List Bullet')
    elif line.startswith('---'):
        # Horizontal rule - skip, add spacing
        doc.add_paragraph()
    elif line.strip() == '':
        # Blank line - skip
        pass
    else:
        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph(line)

    i += 1

# Save
output_path = 'output/liuhan_analysis.docx'
doc.save(output_path)
print(f'Exported to {output_path}')
